from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import re
import shutil
import tempfile
import os

class WordReader:
    def __init__(self, file_path):
        self.file_path = file_path
        self._document = None
        self._temp_file = None

    @property
    def document(self):
        if self._document is None:
            try:
                # Try direct access first
                self._document = Document(self.file_path)
            except PermissionError:
                # If failed, create temp copy
                temp_dir = tempfile.gettempdir()
                self._temp_file = os.path.join(temp_dir, f'temp_{os.path.basename(self.file_path)}')
                shutil.copy2(self.file_path, self._temp_file)
                self._document = Document(self._temp_file)
        return self._document

    def __del__(self):
        # Clean up temp file if it exists
        if self._temp_file and os.path.exists(self._temp_file):
            try:
                os.remove(self._temp_file)
            except:
                pass

    def read_paragraphs(self):
        text = []
        for paragraph in self.document.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)

    def read_tables(self):
        text = []
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return '\n'.join(text)

    def _extract_highlighted_from_paragraph(self, paragraph, color):
        for run in paragraph.runs:
            if run.font.highlight_color and run.font.highlight_color == color:
                yield from self.extract_words(run.text)

    def get_highlighted_words(self, color=WD_COLOR_INDEX.YELLOW):
        for paragraph in self.document.paragraphs:
            yield from self._extract_highlighted_from_paragraph(paragraph, color)
        
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield from self._extract_highlighted_from_paragraph(paragraph, color)

    def extract_words(self, text):
        for word in re.findall(r'\b[a-z]\w*\b', text, re.IGNORECASE):
            yield word

    def get_words(self):
        yield from self.extract_words(self.read_paragraphs())
        yield from self.extract_words(self.read_tables())

if __name__ == '__main__':
    word_reader = WordReader(
        r"C:\Users\Nick\Dropbox\Job Hunting\Applications\In Progress\Data Architect with Warwickshire Police\Data Architect Job Description - October 2024.docx")
    print("\nHighlighted Words:")
    print(list(word_reader.get_highlighted_words()))
    print("\nWords:")
    print(list(word_reader.get_words()))